import re, os, csv, time, yaml, pathlib, requests
from bs4 import BeautifulSoup
from playwright.sync_api import sync_playwright
from docx import Document
from nltk.tokenize import wordpunct_tokenize

ROOT = pathlib.Path(__file__).parent
CFG  = yaml.safe_load(open(ROOT/"config.yaml","r",encoding="utf-8"))
OUT  = ROOT/CFG["resume_engine"]["out_dir"]
OUT.mkdir(parents=True, exist_ok=True)
ART  = ROOT/"artifacts"; ART.mkdir(exist_ok=True)

# ---------- FETCHERS ----------
def fetch_greenhouse(org):
    url = f"https://boards-api.greenhouse.io/v1/boards/{org}/jobs"
    try:
        j = requests.get(url, timeout=20).json().get("jobs", [])
        out=[]
        for x in j:
            out.append({
                "source":"Greenhouse",
                "org": org,
                "title": x.get("title",""),
                "location": ((x.get("location") or {}).get("name","")),
                "url": x.get("absolute_url",""),
                "dept": ((x.get("departments") or [{}])[0]).get("name","")
            })
        return out
    except Exception:
        return []

def fetch_lever(org):
    url = f"https://api.lever.co/v0/postings/{org}?mode=json"
    try:
        posts = requests.get(url, timeout=20).json()
        out=[]
        for p in posts:
            cat = p.get("categories") or {}
            out.append({
                "source":"Lever",
                "org": org,
                "title": p.get("text",""),
                "location": cat.get("location","") or (p.get("workType") or ""),
                "url": p.get("hostedUrl",""),
                "dept": cat.get("team","")
            })
        return out
    except Exception:
        return []

# ---------- FILTERING ----------
def match_job(job, required, exclude):
    t = f"{job['title']} {job['dept']} {job['location']}".lower()
    if not all(re.search(rx.lower(), t) for rx in required): return False
    if any(re.search(rx.lower(), t) for rx in exclude): return False
    return True

def score_job(job):
    t = f"{job['title']} {job['dept']} {job['location']}".lower()
    s=0
    if re.search(r"\banalyst|analytics\b", t): s+=2
    if re.search(r"\bintern|entry|junior\b", t): s+=1
    if re.search(r"ml|machine learning|research|health|bio", t): s+=1
    if re.search(r"remote|hybrid|new york|nyc", t): s+=1
    return -s

# ---------- JD SCRAPER ----------
def jd_text_from_url(url):
    try:
        r = requests.get(url, timeout=20, headers={"User-Agent":"Mozilla/5.0"})
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        container = soup.select_one(".content, .opening, .app-body, .job, .main, main, article") or soup.body
        text = container.get_text("\n", strip=True)
        return re.sub(r"\n{3,}", "\n\n", text)
    except Exception:
        with sync_playwright() as p:
            b = p.chromium.launch(headless=True)
            page = b.new_page()
            page.goto(url, wait_until="domcontentloaded")
            page.wait_for_timeout(3000)
            text = page.inner_text("body")
            b.close()
            return text

# ---------- RESUME ENGINE ----------
def load_yaml(path): return yaml.safe_load(open(path,"r",encoding="utf-8"))

def tokens(text, stopwords):
    toks = [w.lower() for w in wordpunct_tokenize(text)]
    toks = [re.sub(r"[^a-z0-9\+\-#]", "", t) for t in toks]
    return [t for t in toks if t and t not in stopwords]

def tag_score(jd_tokens, tags):
    tags=[t.lower() for t in tags]
    return sum(1 for t in tags if any(t in tok for tok in jd_tokens))

def pick_projects(jd_toks, bank, max_projects=3, bullets_per=2):
    scored=[]
    for key, proj in bank["projects"].items():
        scored.append((tag_score(jd_toks, proj.get("tags",[])), key))
    scored.sort(reverse=True)
    chosen=[]
    for s,key in scored[:max_projects]:
        proj=bank["projects"][key]
        chosen.append({"title":proj["title"], "bullets":proj["bullets"][:bullets_per]})
    gen=[]
    for g in bank.get("general_bullets",[]):
        if tag_score(jd_toks, g.get("tags",[]))>0: gen.append(g["text"])
    return chosen, gen[:2]

def build_resume_doc(base, projects, general, jd_keywords, outfile):
    doc = Document()
    h = doc.add_paragraph()
    r = h.add_run(base["name"]); r.bold=True
    h.add_run(f"  |  {base.get('location','')}  |  {base['phone']}  |  {base['email']}")
    for link in base.get("links", []): h.add_run(f"  |  {link}")
    doc.add_paragraph()
    doc.add_paragraph("SUMMARY").runs[0].bold=True
    doc.add_paragraph(base["summary"])
    doc.add_paragraph(); doc.add_paragraph("SKILLS").runs[0].bold=True
    for line in base.get("skills", []): doc.add_paragraph(line)
    doc.add_paragraph(); doc.add_paragraph("RELEVANT PROJECTS").runs[0].bold=True
    for p in projects:
        pr = doc.add_paragraph(); pr.add_run(p["title"]).bold=True
        for b in p["bullets"]:
            bullet = doc.add_paragraph(b); bullet.style = doc.styles["List Bullet"]
    for gb in general:
        bullet = doc.add_paragraph(gb); bullet.style = doc.styles["List Bullet"]
    doc.add_paragraph(); doc.add_paragraph("EDUCATION").runs[0].bold=True
    for ed in base.get("education", []):
        line = doc.add_paragraph(); line.add_run(ed["school"]).bold=True; line.add_run(f" — {ed['dates']}")
        for b in ed.get("bullets", []):
            bullet = doc.add_paragraph(b); bullet.style = doc.styles["List Bullet"]
    doc.add_paragraph()
    doc.add_paragraph("Keywords: " + ", ".join(sorted(set(jd_keywords)))[:800]).italic=True
    doc.save(outfile)

# ---------- AUTOFILL ----------
def open_and_prefill(url, ident, resume_path, cover_path=None, headless=False, pause_ms=120000):
    with sync_playwright() as p:
        b = p.chromium.launch(headless=headless)
        page = b.new_page()
        page.goto(url, wait_until="domcontentloaded")
        page.wait_for_timeout(1500)

        def try_fill(sel, val):
            try:
                if page.locator(sel).count(): page.fill(sel, val)
            except: pass

        first = ident["name"].split()[0]; last = ident["name"].split()[-1]
        try_fill('input[name*="first"]', first)
        try_fill('input[name*="last"]', last)
        try_fill('input[type="email"]', ident["email"])
        try_fill('input[type="tel"]', ident["phone"])

        for k,v in (ident.get("links") or {}).items():
            try_fill(f'input[placeholder*="{k}"], input[name*="{k}"]', v)

        try:
            file_inputs = page.locator('input[type="file"]').all()
            if file_inputs:
                file_inputs[0].set_input_files(resume_path)
                if cover_path and len(file_inputs) > 1:
                    file_inputs[1].set_input_files(cover_path)
        except: pass

        page.bring_to_front()
        print("✅ Prefilled. Review and click Submit yourself (safer).")
        page.wait_for_timeout(pause_ms)
        b.close()

# ---------- MAIN ----------
def main(top=5):
    ident = CFG["identity"]
    filters = CFG["filters"]

    jobs=[]
    for org in CFG["sources"]["greenhouse"]:
        jobs += fetch_greenhouse(org)
    for org in CFG["sources"]["lever"]:
        jobs += fetch_lever(org)

    req = [r for r in filters["required"]]
    exc = [r for r in filters["exclude"]]
    matches = [j for j in jobs if match_job(j, req, exc)]
    matches.sort(key=score_job)
    Q = matches[:filters.get("top_n", top)]

    base = load_yaml(ROOT/CFG["resume_engine"]["base_resume_file"])
    bank = load_yaml(ROOT/CFG["resume_engine"]["bullets_file"])
    stop = set(w.strip().lower() for w in open(ROOT/CFG["resume_engine"]["stopwords_file"],"r",encoding="utf-8"))
    log_path = ROOT/CFG["log_csv"]
    newlog = not os.path.exists(log_path)

    print(f"\nFound {len(matches)} matching jobs. Taking top {min(len(Q), top)}.\n")
    for i,j in enumerate(Q[:top],1):
        print(f"{i}. {j['title']} — {j['org'] or j['source']}")
        print(f"   {j['url']}\n")

    with open(log_path, "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        if newlog: w.writerow(["time","title","org","url","resume","cover"])
        for j in Q[:top]:
            jd = jd_text_from_url(j["url"])
            jd_toks = [t for t in tokens(jd, stop) if t]
            projects, general = pick_projects(jd_toks, bank, max_projects=3, bullets_per=2)

            ts = int(time.time())
            resume_out = OUT/f"resume_{ts}.docx"
            cover_out  = OUT/f"cover_{ts}.docx"

            build_resume_doc(base, projects, general, jd_toks, str(resume_out))

            doc = Document()
            doc.add_paragraph("Dear Hiring Team,")
            doc.add_paragraph(f"I’m applying for {j['title']}. My experience in data/ML and healthcare includes projects like "
                              f"{projects[0]['title'] if projects else 'NLP matching'} and PySpark analytics.")
            doc.add_paragraph("I value clean, reproducible pipelines and clear communication. Excited to contribute.")
            doc.add_paragraph("\nThank you,\n" + base["name"])
            doc.save(str(cover_out))

            try:
                open_and_prefill(j["url"], CFG["identity"], str(resume_out), str(cover_out),
                                 headless=not (not CFG["autofill"]["headless"]),
                                 pause_ms=int(CFG["autofill"]["human_review_ms"]))
            except Exception as e:
                print("Prefill issue (skipping to next):", e)

            w.writerow([time.time(), j["title"], j.get("org",""), j["url"], str(resume_out), str(cover_out)])

if __name__ == "__main__":
    main(top=5)
