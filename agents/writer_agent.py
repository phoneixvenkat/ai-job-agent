import os
import json
from dotenv import load_dotenv
from langchain_groq import ChatGroq
import time
import pathlib
import yaml
from docx import Document
from langchain_core.messages import HumanMessage
import urllib3
from backend.utils.logger import get_logger
log = get_logger('writer')

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

ROOT = pathlib.Path(__file__).parent.parent
OUT  = ROOT / "out"
OUT.mkdir(exist_ok=True)

load_dotenv()
_llm = None

def _get_llm():
    global _llm
    if _llm is None:
        _llm = ChatGroq(model=os.getenv("GROQ_MODEL", "llama-3.1-8b-instant"), api_key=os.getenv("GROQ_API_KEY"), temperature=0.3)
    return _llm

def load_yaml(path):
    return yaml.safe_load(open(path, "r", encoding="utf-8"))

def tag_score(jd_tokens, tags):
    tags = [t.lower() for t in tags]
    return sum(1 for t in tags if any(t in tok for tok in jd_tokens))

def pick_bullets(jd_text: str, bank: dict, max_projects=3, bullets_per=2) -> tuple:
    import re
    from nltk.tokenize import wordpunct_tokenize
    try:
        import nltk
        nltk.download("punkt", quiet=True)
        nltk.download("punkt_tab", quiet=True)
    except: pass

    stop    = set(w.strip().lower() for w in open(ROOT/"data"/"stopwords.txt", "r", encoding="utf-8"))
    toks    = [re.sub(r"[^a-z0-9\+\-#]", "", w.lower()) for w in wordpunct_tokenize(jd_text)]
    jd_toks = [t for t in toks if t and t not in stop]

    scored = []
    for key, proj in bank.get("projects", {}).items():
        scored.append((tag_score(jd_toks, proj.get("tags", [])), key))
    scored.sort(reverse=True)

    chosen = []
    for _, key in scored[:max_projects]:
        proj = bank["projects"][key]
        chosen.append({"title": proj["title"], "bullets": proj["bullets"][:bullets_per]})

    general = []
    for g in bank.get("general_bullets", []):
        if tag_score(jd_toks, g.get("tags", [])) > 0:
            general.append(g["text"])

    return chosen, general[:2], jd_toks


def extract_experience_llm(resume_text: str, job_title: str, jd_text: str) -> list:
    """Parse the uploaded resume text and return structured work/internship experience."""
    prompt = f"""Extract all work experience, internships, and research positions from this resume text.
For each role return a JSON object with: title, company, dates, bullets (2-3 achievement bullets).
Prioritize entries most relevant to: {job_title}
Job context: {jd_text[:300]}

Resume:
{resume_text[:3000]}

Return ONLY a valid JSON array. If no experience found, return [].
Example format:
[
  {{"title": "Data Science Intern", "company": "Acme Corp", "dates": "Jun 2023 – Aug 2023",
    "bullets": ["Built ML pipeline improving accuracy by 15%", "Processed 1M+ records using PySpark"]}}
]"""
    try:
        response = _get_llm().invoke([HumanMessage(content=prompt)])
        text = response.content.strip()
        start = text.find('[')
        end   = text.rfind(']') + 1
        if start >= 0 and end > start:
            return json.loads(text[start:end])
    except Exception:
        pass
    return []


def rewrite_summary_llm(base_summary: str, job_title: str, jd_text: str, company: str = "") -> str:
    company_line = f"Company: {company}\n" if company else ""
    prompt = f"""Rewrite this professional summary to better match the job.
Keep all facts accurate. Maximum 3 sentences. Sound human and natural.
{"Mention " + company + " by name in the summary." if company else ""}

Job Title: {job_title}
{company_line}Job Description (first 500 chars): {jd_text[:500]}
Original Summary: {base_summary}

Return only the rewritten summary, nothing else."""
    try:
        response = _get_llm().invoke([HumanMessage(content=prompt)])
        return response.content.strip()
    except:
        return base_summary


def generate_cover_letter_llm(base: dict, job: dict, jd_text: str) -> str:
    prompt = f"""Write a professional, personalized cover letter for this job application.

Candidate: {base['name']}
Job Title: {job['title']}
Company: {job['org']}
Job Description (first 600 chars): {jd_text[:600]}
Candidate Background: {base.get('summary', '')}
Key Skills: {', '.join(base.get('skills', [])[:5]) if base.get('skills') else 'Python, ML, Data Science'}

Instructions:
- 3 short paragraphs
- Sound human and enthusiastic
- Mention 2-3 specific skills matching the role
- End with a call to action
- Do NOT use placeholders like [Your Name]
- Return only the letter text"""
    try:
        response = _get_llm().invoke([HumanMessage(content=prompt)])
        return response.content.strip()
    except:
        return f"Dear Hiring Team at {job['org']},\n\nI am excited to apply for the {job['title']} position.\n\nThank you,\n{base['name']}"


def build_resume_doc(base: dict, projects: list, general: list, experience: list, summary: str, outfile: str):
    doc = Document()
    h = doc.add_paragraph()
    r = h.add_run(base["name"]); r.bold = True
    h.add_run(f"  |  {base.get('location','')}  |  {base['phone']}  |  {base['email']}")
    for link in base.get("links", []):
        h.add_run(f"  |  {link}")

    doc.add_paragraph()
    doc.add_paragraph("SUMMARY").runs[0].bold = True
    doc.add_paragraph(summary)

    doc.add_paragraph()
    doc.add_paragraph("SKILLS").runs[0].bold = True
    for line in base.get("skills", []):
        doc.add_paragraph(line)

    if experience:
        doc.add_paragraph()
        doc.add_paragraph("EXPERIENCE").runs[0].bold = True
        for exp in experience:
            line = doc.add_paragraph()
            line.add_run(exp.get("title", "")).bold = True
            if exp.get("company"):
                line.add_run(f" — {exp['company']}")
            if exp.get("dates"):
                line.add_run(f"  |  {exp['dates']}")
            for b in exp.get("bullets", []):
                bullet = doc.add_paragraph(b)
                bullet.style = doc.styles["List Bullet"]

    doc.add_paragraph()
    doc.add_paragraph("RELEVANT PROJECTS").runs[0].bold = True
    for p in projects:
        pr = doc.add_paragraph()
        pr.add_run(p["title"]).bold = True
        for b in p["bullets"]:
            bullet = doc.add_paragraph(b)
            bullet.style = doc.styles["List Bullet"]
    for gb in general:
        bullet = doc.add_paragraph(gb)
        bullet.style = doc.styles["List Bullet"]

    doc.add_paragraph()
    doc.add_paragraph("EDUCATION").runs[0].bold = True
    for ed in base.get("education", []):
        line = doc.add_paragraph()
        line.add_run(ed["school"]).bold = True
        line.add_run(f" — {ed['dates']}")
        for b in ed.get("bullets", []):
            bullet = doc.add_paragraph(b)
            bullet.style = doc.styles["List Bullet"]

    doc.save(outfile)


def run_writer(job: dict, resume_text: str, use_llm: bool = True) -> dict:
    log.info(f"\n  Writer Agent: {job['title']} at {job['org']}")

    base = load_yaml(ROOT/"data"/"base_resume.yaml")
    bank = load_yaml(ROOT/"data"/"bullet_bank.yaml")
    jd   = job.get("description", job.get("title", ""))

    projects, general, _ = pick_bullets(jd, bank)

    # Extract experience from uploaded resume
    experience = []
    if use_llm and resume_text.strip():
        log.info("   Extracting experience from uploaded resume...")
        experience = extract_experience_llm(resume_text, job["title"], jd)
        log.info(f"   Found {len(experience)} experience entries")

    # LLM rewrite summary with company context
    if use_llm:
        log.info("   Rewriting summary with LLM...")
        summary = rewrite_summary_llm(base.get("summary", ""), job["title"], jd, company=job.get("org", ""))
    else:
        summary = base.get("summary", "")

    # Build resume DOCX
    ts         = f"{job['org']}_{int(time.time())}".replace(" ", "_")
    resume_out = str(OUT / f"resume_{ts}.docx")
    build_resume_doc(base, projects, general, experience, summary, resume_out)
    log.info(f"   Resume saved: resume_{ts}.docx")

    # Generate cover letter
    if use_llm:
        log.info("   Writing cover letter with LLM...")
        cover_text = generate_cover_letter_llm(base, job, jd)
    else:
        cover_text = f"Dear Hiring Team at {job['org']},\n\nI am applying for {job['title']}.\n\nThank you,\n{base['name']}"

    cover_doc = Document()
    cover_doc.add_paragraph("Dear Hiring Team,")
    for para in cover_text.split("\n"):
        if para.strip():
            cover_doc.add_paragraph(para)
    cover_out = str(OUT / f"cover_{ts}.docx")
    cover_doc.save(cover_out)
    log.info(f"   Cover letter saved: cover_{ts}.docx")

    return {
        "job":          job,
        "resume_path":  resume_out,
        "cover_path":   cover_out,
        "cover_text":   cover_text,
        "summary_used": summary,
        # Structured sections for frontend preview
        "name":         base["name"],
        "contact":      f"{base.get('location','')} | {base['email']} | {base['phone']}",
        "skills":       base.get("skills", []),
        "projects":     projects,
        "experience":   experience,
        "education":    base.get("education", []),
    }
